from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


BASE_DIR = Path(r"C:\Users\moham\Downloads")
WORKSPACE = Path(__file__).resolve().parents[1]
OUTPUT_DIR = WORKSPACE / "output" / "doc"

SRC_AI_ML = BASE_DIR / "AI ML Book.docx"
SRC_QBANK = BASE_DIR / "Part 13 Interview Question Bank.docx"


TITLE_BLUE = RGBColor(31, 77, 120)
TITLE_DARK = RGBColor(11, 37, 69)
BODY_DARK = RGBColor(51, 51, 51)
MUTED = RGBColor(90, 90, 90)


class Block:
    def __init__(self, kind: str, element, text: str = "", style: str = ""):
        self.kind = kind
        self.element = element
        self.text = text
        self.style = style


def parse_blocks(doc: Document) -> list[Block]:
    blocks: list[Block] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            blocks.append(Block("p", child, para.text, para.style.name if para.style else ""))
        elif child.tag == qn("w:tbl"):
            blocks.append(Block("tbl", child))
    return blocks


def set_run_font(run, *, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    r = run._element.get_or_add_rPr()
    rFonts = r.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_DARK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after, bold, align, page_break in [
        ("Heading 1", 18, TITLE_DARK, 20, 6, True, WD_ALIGN_PARAGRAPH.CENTER, True),
        ("Heading 2", 14, TITLE_BLUE, 18, 6, True, WD_ALIGN_PARAGRAPH.LEFT, True),
        ("Heading 3", 12, TITLE_DARK, 12, 4, True, WD_ALIGN_PARAGRAPH.LEFT, False),
        ("Heading 4", 11, BODY_DARK, 8, 3, True, WD_ALIGN_PARAGRAPH.LEFT, False),
        ("Heading 5", 10, MUTED, 6, 2, True, WD_ALIGN_PARAGRAPH.LEFT, False),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15 if name in {"Heading 1", "Heading 2"} else 1.1
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.page_break_before = page_break
        st.paragraph_format.alignment = align

    lp = doc.styles["List Paragraph"]
    lp.font.name = "Calibri"
    lp.font.size = Pt(11)
    lp.font.color.rgb = BODY_DARK
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(3)
    lp.paragraph_format.line_spacing = 1.15


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break()


def add_centered_paragraph(doc: Document, text: str, size: int, *, color=BODY_DARK, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_left_paragraph(doc: Document, text: str, size: int = 11, *, color=BODY_DARK, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_title_page(doc: Document, volume_title: str, subtitle: str) -> None:
    add_centered_paragraph(doc, "AI/ML Engineer Internship & Fresher Preparation Handbook", 12, color=MUTED, bold=True, after=14)
    add_centered_paragraph(doc, volume_title, 26, color=TITLE_DARK, bold=True, after=10)
    add_centered_paragraph(doc, subtitle, 16, color=TITLE_BLUE, bold=False, after=14)
    add_centered_paragraph(doc, "Professional multi-volume technical book series", 11, color=MUTED, italic=True, after=26)
    add_centered_paragraph(doc, "Prepared from the supplied source documents", 10, color=MUTED, after=0)


def add_simple_section(doc: Document, heading: str, paragraphs: Sequence[str]) -> None:
    doc.add_paragraph(heading, style="Heading 1")
    for text in paragraphs:
        add_left_paragraph(doc, text, after=6)


def add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("• ")
        set_run_font(r, size=11, color=BODY_DARK)
        r2 = p.add_run(item)
        set_run_font(r2, size=11, color=BODY_DARK)


def add_toc_table(doc: Document, rows: Sequence[Tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.2), Inches(5.8)]
    hdr = table.rows[0].cells
    hdr[0].width = widths[0]
    hdr[1].width = widths[1]
    hdr[0].text = "Chapter"
    hdr[1].text = "Title"
    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=10.5, color=TITLE_DARK, bold=True)
    for chapter, title in rows:
        cells = table.add_row().cells
        cells[0].width = widths[0]
        cells[1].width = widths[1]
        cells[0].text = chapter
        cells[1].text = title
        for p in cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=10.5, color=BODY_DARK)
        for p in cells[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                set_run_font(r, size=10.5, color=BODY_DARK)


def add_index(doc: Document, entries: Sequence[str]) -> None:
    doc.add_paragraph("Index", style="Heading 1")
    add_bullets(doc, entries)


def is_part_heading(text: str) -> bool:
    return text.startswith("Part ")


def is_chapter_heading(text: str) -> bool:
    return text.startswith("Chapter ")


def strip_heading_prefix(text: str) -> str:
    m = re.match(r"^Chapter\s+[0-9.]+[:\-\s]*(.*)$", text)
    if m:
        tail = m.group(1).strip()
        return tail if tail else text
    m = re.match(r"^Chapter\s+([0-9.]+)\s*(.*)$", text)
    if m:
        tail = m.group(2).strip()
        if tail.startswith(":"):
            tail = tail[1:].strip()
        return tail if tail else text
    return text


def set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    paragraph.text = new_text


def copy_blocks(
    src_blocks: Sequence[Block],
    dst: Document,
    *,
    chapter_style_map: dict,
    volume_key: str,
    toc_rows: list[Tuple[str, str]],
    index_entries: list[str],
) -> None:
    chapter_count = 0
    for block in src_blocks:
        if block.kind == "p":
            para = Paragraph(block.element, dst)
            text = para.text.strip()
            style = para.style.name if para.style else ""

            if is_part_heading(text):
                new_el = deepcopy(block.element)
                body = dst.element.body
                body.insert(len(body) - 1, new_el)
                p = Paragraph(new_el, dst)
                p.style = dst.styles["Heading 1"]
                p.paragraph_format.page_break_before = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if volume_key == "qbank" and text.startswith("Part 14"):
                    pass
                continue

            if is_chapter_heading(text):
                chapter_count += 1
                title = strip_heading_prefix(text)
                chapter_label = f"Chapter {chapter_count}"
                new_text = f"{chapter_label} - {title}" if title else chapter_label
                toc_rows.append((chapter_label, title or text))
                index_entries.append(title or text)
                new_el = deepcopy(block.element)
                body = dst.element.body
                body.insert(len(body) - 1, new_el)
                p = Paragraph(new_el, dst)
                p.style = dst.styles["Heading 2"]
                p.paragraph_format.page_break_before = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_text(p, new_text)
                continue

            if volume_key == "qbank" and style == "Heading 1":
                # Demote question and major topic headings so the question bank reads like a book.
                new_el = deepcopy(block.element)
                body = dst.element.body
                body.insert(len(body) - 1, new_el)
                p = Paragraph(new_el, dst)
                if text.startswith("Part "):
                    p.style = dst.styles["Heading 1"]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.page_break_before = True
                else:
                    p.style = dst.styles["Heading 3"]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            if volume_key == "qbank" and style == "Heading 2":
                new_el = deepcopy(block.element)
                body = dst.element.body
                body.insert(len(body) - 1, new_el)
                p = Paragraph(new_el, dst)
                p.style = dst.styles["Heading 4"]
                continue

            new_el = deepcopy(block.element)
            body = dst.element.body
            body.insert(len(body) - 1, new_el)
        else:
            new_el = deepcopy(block.element)
            body = dst.element.body
            body.insert(len(body) - 1, new_el)


def make_volume(
    volume_title: str,
    subtitle: str,
    src_blocks: Sequence[Block],
    volume_key: str,
    toc_rows: list[Tuple[str, str]],
    index_entries: list[str],
    learning_points: Sequence[str],
    preface_text: Sequence[str],
    how_to_text: Sequence[str],
) -> Document:
    dst = Document()
    configure_styles(dst)
    add_title_page(dst, volume_title, subtitle)
    dst.add_page_break()
    add_simple_section(dst, "Copyright Page", [
        "Copyright 2026. Prepared from the user-supplied source documents for editorial repackaging and publication-style formatting.",
        "All educational content has been retained or reorganized for clarity, consistency, and navigation.",
    ])
    dst.add_page_break()
    add_simple_section(dst, "Preface", preface_text)
    dst.add_page_break()
    add_simple_section(dst, "How to Use This Volume", how_to_text)
    dst.add_page_break()
    add_simple_section(dst, "Learning Outcomes", [])
    add_bullets(dst, learning_points)
    dst.add_page_break()
    doc_toc_rows: list[Tuple[str, str]] = []
    add_simple_section(dst, "Complete Table of Contents", [])
    add_toc_table(dst, toc_rows)
    dst.add_page_break()
    copy_blocks(src_blocks, dst, chapter_style_map={}, volume_key=volume_key, toc_rows=doc_toc_rows, index_entries=index_entries)
    dst.add_page_break()
    add_simple_section(dst, "Revision Notes", [
        "This volume preserves the source material while presenting it in a cleaner chapter sequence and a more book-like layout.",
        "Use the chapter headings and quick revision sections as the main review path before interviews or assessment sessions.",
    ])
    add_index(dst, index_entries)
    return dst


def collect_chapter_rows(src_blocks: Sequence[Block], volume_key: str) -> list[Tuple[str, str]]:
    rows: list[Tuple[str, str]] = []
    chapter_count = 0
    for block in src_blocks:
        if block.kind != "p":
            continue
        text = block.text.strip()
        if is_chapter_heading(text):
            chapter_count += 1
            rows.append((f"Chapter {chapter_count}", strip_heading_prefix(text) or text))
    return rows


def select_ai_ml_volume_blocks(all_blocks: Sequence[Block], start_part: int, end_part: int | None) -> list[Block]:
    selected: list[Block] = []
    part_re = re.compile(r"^Part\s+(\d+):")
    current = None
    for block in all_blocks:
        if block.kind == "p":
            text = block.text.strip()
            m = part_re.match(text)
            if m:
                part_num = int(m.group(1))
                current = part_num
            if current is not None and current >= start_part and (end_part is None or current <= end_part):
                selected.append(block)
        elif current is not None and current >= start_part and (end_part is None or current <= end_part):
            selected.append(block)
    return selected


def build_master_toc(volume_specs: Sequence[Tuple[str, str, str, str, str]]) -> Document:
    doc = Document()
    configure_styles(doc)
    add_title_page(doc, "Master Series Table of Contents", "AI/ML Engineer Internship & Fresher Preparation Handbook")
    doc.add_page_break()
    add_simple_section(doc, "Series Overview", [
        "This series repackages the supplied handbook material into five publication-style volumes with cleaner chapter numbering, clearer front matter, and a more book-like reading order.",
        "Each volume is designed to stand on its own while also fitting into the larger AI/ML interview-preparation pathway.",
    ])
    doc.add_page_break()
    doc.add_paragraph("Volume Map", style="Heading 1")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Volume", "Title", "Reading Order", "Estimated Pages"]
    widths = [Inches(0.8), Inches(3.2), Inches(1.5), Inches(1.0)]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=10.5, color=TITLE_DARK, bold=True)
    for vol, title, order, estimate, _ in volume_specs:
        row = table.add_row().cells
        vals = [vol, title, order, estimate]
        for i, val in enumerate(vals):
            row[i].width = widths[i]
            row[i].text = val
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i != 0 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_run_font(r, size=10.5, color=BODY_DARK)
    doc.add_page_break()
    add_simple_section(doc, "Recommended Reading Order", [
        "1. Volume 1 - Python, Mathematics, and core machine learning foundations",
        "2. Volume 2 - Deep learning, computer vision, and natural language processing",
        "3. Volume 3 - Generative AI, MLOps, deployment, and SQL",
        "4. Volume 4 - Project development and project discussion mastery",
        "5. Volume 5 - Final interview preparation, revision sheets, and mock interviews",
    ])
    doc.add_page_break()
    add_simple_section(doc, "Learning Path Overview", [
        "Build the base first: Python, mathematics, and classical machine learning.",
        "Move into deep learning and applied domain work: computer vision and NLP.",
        "Add production skills: generative AI, MLOps, deployment, and SQL.",
        "Use the project and interview volumes to turn knowledge into confident explanations.",
        "Finish with timed revision, question banks, and mock interview practice.",
    ])
    return doc


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> None:
    if not SRC_AI_ML.exists():
        raise FileNotFoundError(SRC_AI_ML)
    if not SRC_QBANK.exists():
        raise FileNotFoundError(SRC_QBANK)

    ai_doc = Document(str(SRC_AI_ML))
    qbank_doc = Document(str(SRC_QBANK))

    ai_blocks = parse_blocks(ai_doc)
    qbank_blocks = parse_blocks(qbank_doc)

    volume1_blocks = select_ai_ml_volume_blocks(ai_blocks, 1, 5)
    volume2_blocks = select_ai_ml_volume_blocks(ai_blocks, 6, 8)
    volume3_blocks = select_ai_ml_volume_blocks(ai_blocks, 9, 11)
    volume4_blocks = select_ai_ml_volume_blocks(ai_blocks, 12, 12)
    volume5_blocks = list(qbank_blocks)

    volume1_toc = collect_chapter_rows(volume1_blocks, "ai")
    volume2_toc = collect_chapter_rows(volume2_blocks, "ai")
    volume3_toc = collect_chapter_rows(volume3_blocks, "ai")
    volume4_toc = collect_chapter_rows(volume4_blocks, "ai")
    volume5_toc = collect_chapter_rows(volume5_blocks, "qbank")

    series_specs = [
        ("Volume 1", "Python, Mathematics & Machine Learning Foundations", "Foundations", "70-90", "Volume 1"),
        ("Volume 2", "Deep Learning, Computer Vision & NLP", "Applied AI", "65-85", "Volume 2"),
        ("Volume 3", "Generative AI, MLOps & SQL", "Production Skills", "70-95", "Volume 3"),
        ("Volume 4", "Project Development & Project Discussion Mastery", "Projects", "45-70", "Volume 4"),
        ("Volume 5", "Final Interview Preparation Guide", "Interview Prep", "90-130", "Volume 5"),
    ]

    vol1 = make_volume(
        "Volume 1",
        "Python, Mathematics & Machine Learning Foundations",
        volume1_blocks,
        "ai",
        volume1_toc,
        ["Python", "NumPy", "Pandas", "Linear Algebra", "Statistics", "Machine Learning Fundamentals", "Model Evaluation"],
        [
            "Use this volume to build the core programming and analytical base for AI/ML interviews.",
            "Read the chapters in order if you want a structured learning path, or jump directly to a topic if you already know the basics.",
            "The chapter numbering has been normalized so the volume reads like a finished book rather than lecture notes.",
        ],
        [
            "Start with Python and data structures, then work through mathematics and classical machine learning.",
            "Treat the algorithm chapters as reference material and revisit the evaluation chapter before interviews.",
            "Use the table of contents as a reading map and the index to revisit fast-changing concepts.",
        ],
        [
            "Move from the basics upward in the same order as the contents list so the foundation stays coherent.",
            "Pause after each major part and review the chapter headings before moving on.",
            "Use the revision notes at the end of the volume as a final checkpoint.",
        ],
    )
    vol2 = make_volume(
        "Volume 2",
        "Deep Learning, Computer Vision & NLP",
        volume2_blocks,
        "ai",
        volume2_toc,
        ["Deep Learning", "Neural Networks", "CNNs", "Computer Vision", "NLP", "Transformers", "BERT", "GPT"],
        [
            "Use this volume to connect foundational machine learning with the domain skills that frequently appear in AI engineer interviews.",
            "The chapter sequence is arranged to move from neural-network basics into vision and language workflows.",
            "Tables, lists, and explanatory sections from the source material are preserved to keep the original teaching depth intact.",
        ],
        [
            "Read the deep learning chapters first so CNNs and transformers make more sense in later sections.",
            "Use the computer vision and NLP chapters as applied references for project explanations and interview follow-ups.",
            "Review the quick revision material before mock interviews or written tests.",
        ],
        [
            "Work through the domain sections in order, then revisit the revision material as a fast interview refresher.",
            "If you are short on time, focus on chapter titles and revision pages first, then return to the detailed explanations.",
            "Use the chapter order as a ready-made study plan.",
        ],
    )
    vol3 = make_volume(
        "Volume 3",
        "Generative AI, MLOps & SQL",
        volume3_blocks,
        "ai",
        volume3_toc,
        ["Generative AI", "LLMs", "Prompt Engineering", "RAG", "MLOps", "Docker", "APIs", "SQL"],
        [
            "This volume bridges model understanding with production workflows and database fundamentals.",
            "The chapters are organized so the reader can move from generative AI concepts into the tooling used to ship and monitor models.",
            "SQL is included here as a practical interview skill and a working companion to data and deployment topics.",
        ],
        [
            "Read the generative AI chapters alongside the MLOps section to understand how modern AI systems are built and deployed.",
            "Treat SQL as a recurring interview skill and practice it regularly instead of reading it once.",
            "Use the revision sections to sharpen concise explanations for interviews and project discussions.",
        ],
        [
            "Use the first half of the volume for concepts and the second half for toolchain and database practice.",
            "Review the quick revision sheets after each major topic area so the terminology stays fresh.",
            "Make SQL practice a repeated habit rather than a one-time reading pass.",
        ],
    )
    vol4 = make_volume(
        "Volume 4",
        "Project Development & Project Discussion Mastery",
        volume4_blocks,
        "ai",
        volume4_toc,
        ["Project Explanation", "Documentation", "STAR Method", "Follow-Up Questions", "Portfolio Building", "Case Studies", "Discussion Skills"],
        [
            "This volume turns technical knowledge into interview-ready project storytelling.",
            "It preserves the source guidance on how to explain projects, answer follow-up questions, and communicate outcomes clearly.",
            "Use the material to shape both your portfolio narrative and your in-person discussion style.",
        ],
        [
            "Practice explaining one project in a 60-second version and a 2-minute version.",
            "Rehearse follow-up questions until you can answer them without sounding memorized.",
            "Use the case-study style chapters to improve structure, confidence, and technical clarity.",
        ],
        [
            "Treat the project discussion chapters as rehearsal material and say the answers out loud.",
            "Use the section order to build a crisp project story with problem, method, result, and lessons learned.",
            "Return to the follow-up questions after you refine your project narrative.",
        ],
    )
    vol5 = make_volume(
        "Volume 5",
        "Final Interview Preparation Guide",
        volume5_blocks,
        "qbank",
        volume5_toc,
        ["Deep Learning Questions", "SQL Questions", "AI/ML Internship Questions", "HR Questions", "Revision Sheets", "Cheat Sheets", "Mock Interviews"],
        [
            "This volume is the interview-facing capstone of the series and preserves the detailed question-and-answer material from the second source document.",
            "The question bank, revision sheets, mock interview, and HR preparation sections are presented in a more polished book format.",
            "Use it as the final review layer after working through the technical foundation volumes.",
        ],
        [
            "Work through the deep learning and SQL question sets first, then move into the larger AI/ML interview bank and HR practice.",
            "Use the quick revision sheets for timed review sessions and last-minute memorization.",
            "Treat the mock interview and final advice sections as a rehearsal script before real interviews.",
        ],
        [
            "Read this volume after the foundations so the questions feel like reinforcement rather than a first exposure.",
            "Use the revision sheets to compress the whole handbook into one last pass before interviews.",
            "Keep the mock interview section open while you practice answers aloud.",
        ],
    )

    save_doc(vol1, OUTPUT_DIR / "Volume_1_Python_Math_ML_Foundations.docx")
    save_doc(vol2, OUTPUT_DIR / "Volume_2_DeepLearning_ComputerVision_NLP.docx")
    save_doc(vol3, OUTPUT_DIR / "Volume_3_GenerativeAI_MLOps_SQL.docx")
    save_doc(vol4, OUTPUT_DIR / "Volume_4_Project_Development_and_Discussion.docx")
    save_doc(vol5, OUTPUT_DIR / "Volume_5_Final_Interview_Preparation_Guide.docx")

    master = build_master_toc(series_specs)
    save_doc(master, OUTPUT_DIR / "Master_Series_Table_of_Contents.docx")

    print(f"Created documents in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
